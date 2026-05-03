"""
generate_docx.py
=============================================================================
Generates Method_Statement_RCC.docx from generated_sections_debug.json

Usage:
    python generate_docx.py \
        --sections_json output/generated_sections_debug.json \
        --out_path   output/Method_Statement_RCC.docx \
        --team_name  "Your Team Name"
"""

import os
import re
import json
import argparse
from datetime import datetime

# ── Word output ───────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ══════════════════════════════════════════════════════════════════════════
# SECTION DEFINITIONS (must match retrieve_and_generate.py)
# ══════════════════════════════════════════════════════════════════════════
MS_SECTIONS = [
    {"key": "purpose", "heading": "1. Purpose of the Method Statement"},
    {"key": "scope", "heading": "2. Scope of the Method Statement"},
    {"key": "acronyms", "heading": "3. Acronyms and Definitions"},
    {"key": "references", "heading": "4. Reference Documents"},
    {"key": "procedure", "heading": "5. Procedure for Concreting"},
    {"key": "equipment", "heading": "6. Equipment Used"},
    {"key": "personnel", "heading": "7. Key People Involved"},
    {"key": "quality", "heading": "8. Quality Control and Testing"},
    {"key": "health_safety", "heading": "9. Health, Safety & Environment (HSE) Considerations"},
    {"key": "other", "heading": "10. Other Relevant Information"},
]


# ══════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════

def build_word_document(sections_content: dict[str, str],
                        team_name: str,
                        out_path: str) -> None:
    """Generate Word document from sections content dictionary."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    title = doc.add_heading("METHOD STATEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading("Reinforced Cement Concrete (RCC) Works", level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: {team_name}\n").bold = True
    meta.add_run(f"Date: {datetime.today().strftime('%d %B %Y')}\n")
    meta.add_run(
        "Document Reference: Based on CPWD Prescriptive Specifications\n"
    )

    doc.add_page_break()

    for sec in MS_SECTIONS:
        key     = sec["key"]
        heading = sec["heading"]
        content = sections_content.get(key, "Content not generated.")

        doc.add_heading(heading, level=1)

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if re.match(r"^[-*•]\s+", stripped):
                p = doc.add_paragraph(
                    stripped.lstrip("-*• "), style="List Bullet"
                )
            elif re.match(r"^\d+[\.\)]\s+", stripped):
                p = doc.add_paragraph(
                    re.sub(r"^\d+[\.\)]\s+", "", stripped),
                    style="List Number",
                )
            else:
                p = doc.add_paragraph(stripped)
            if p.runs:
                p.runs[0].font.size = Pt(11)

        doc.add_paragraph()

    doc.add_heading("Document Conventions", level=2)
    doc.add_paragraph(
        "Throughout this document, the symbol '§' (section sign) refers to specific "
        "clauses or subsections within the CPWD Prescriptive Specifications. "
        "For example, '§5.1.3' indicates Clause 5.1.3 of the specification document."
    )

    doc.add_heading("Note on Sources", level=2)
    doc.add_paragraph(
        "All information in this Method Statement has been extracted from the "
        "CPWD Prescriptive Specifications document using an NLP-based retrieval "
        "pipeline (parser v3 + embed_and_store v4).  Where specification text "
        "was insufficient, this has been explicitly noted."
    )

    doc.save(out_path)
    print(f"✅  Word document saved → {out_path}")


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

def main(sections_json: str, out_path: str, team_name: str = "Your Team Name") -> None:
    """Read sections JSON and generate docx."""
    if not os.path.exists(sections_json):
        raise FileNotFoundError(f"Sections JSON not found: {sections_json}")

    print(f"📖  Loading sections from {sections_json} …")
    with open(sections_json, "r", encoding="utf-8") as f:
        sections_content = json.load(f)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    build_word_document(sections_content, team_name, out_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate Method Statement docx from generated_sections_debug.json"
    )
    parser.add_argument("--sections_json", default="output/generated_sections_debug.json",
                        help="Path to generated_sections_debug.json")
    parser.add_argument("--out_path", default="output/Method_Statement_RCC.docx",
                        help="Path for output docx file")
    parser.add_argument("--team_name", default="musketeers",
                        help="Team name for the title page")
    args = parser.parse_args()

    main(args.sections_json, args.out_path, args.team_name)
