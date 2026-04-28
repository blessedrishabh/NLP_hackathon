"""
retrieve_and_generate.py
========================
Stage 4 — Retrieval  : targeted ChromaDB queries per Method-Statement section
Stage 5 — LLM        : Gemini 1.5 Flash (free) writes each section from retrieved chunks

Usage (Google Colab / local):
    pip install chromadb sentence-transformers google-generativeai python-docx

    python retrieve_and_generate.py \
        --db_dir   chroma_db \
        --out_dir  output \
        --api_key  YOUR_GEMINI_API_KEY        # or set env GEMINI_API_KEY
"""

import os
import json
import re
import argparse
from datetime import datetime

# ── vector store + embeddings ──────────────────────────────────────────────
import chromadb
from sentence_transformers import SentenceTransformer

# ── LLM ───────────────────────────────────────────────────────────────────
import google.generativeai as genai

# ── Word output ───────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from sentence_transformers import CrossEncoder
load_dotenv()


# ══════════════════════════════════════════════════════════════════════════
# 1.  SECTION DEFINITIONS
#     Each entry has:
#       key        – internal identifier
#       heading    – heading text in the Word output
#       queries    – list of retrieval queries (results are pooled & de-duped)
#       word_limit – soft cap passed to the LLM
#       instruction– extra guidance given to the LLM for that section
# ══════════════════════════════════════════════════════════════════════════
MS_SECTIONS = [
    {
        "key": "purpose",
        "heading": "1. Purpose of the Method Statement",
        "queries": [
            "purpose scope objective of reinforced cement concrete RCC work",
            "general requirements for RCC construction specification",
        ],
        "word_limit": 120,
        "instruction": (
            "Write a concise purpose statement (2–3 sentences) explaining WHY "
            "this method statement exists, referencing the RCC works covered by "
            "the specification."
        ),
    },
    {
        "key": "scope",
        "heading": "2. Scope of the Method Statement",
        "queries": [
            "scope of work reinforced cement concrete structural elements",
            "applicability RCC specification foundations columns beams slabs",
        ],
        "word_limit": 150,
        "instruction": (
            "Define WHAT work is covered — structural elements, locations, "
            "grade of concrete — citing the specification sections."
        ),
    },
    {
        "key": "acronyms",
        "heading": "3. Acronyms and Definitions",
        "queries": [
            "abbreviations acronyms definitions terminology concrete specification",
            "IS code BIS OPC PPC w/c ratio slump workability definitions",
        ],
        "word_limit": 200,
        "instruction": (
            "List all acronyms and technical terms found in the retrieved text "
            "as a definition list (TERM – meaning).  Do NOT invent definitions "
            "not present in the specification."
        ),
    },
    {
        "key": "references",
        "heading": "4. Reference Documents",
        "queries": [
            "IS code standards references BIS bureau Indian standard concrete",
            "relevant codes specifications standards cited in document",
        ],
        "word_limit": 200,
        "instruction": (
            "List every IS code, BIS standard, or other document explicitly "
            "cited in the retrieved chunks.  Format as a numbered list."
        ),
    },
    {
        "key": "procedure",
        "heading": "5. Procedure for Concreting",
        "queries": [
            "procedure batching mixing transporting placing compacting curing concrete",
            "concrete mix design proportioning water cement ratio aggregate",
            "formwork shuttering reinforcement bar bending placing cover",
            "compaction vibration consolidation concrete placing procedure",
            "curing methods period water curing membrane curing concrete",
            "concrete pour sequence joints construction cold joints",
            "admixtures plasticizer retarder accelerator concrete mixing",
        ],
        "word_limit": 600,
        "instruction": (
            "Write a step-by-step concreting procedure with numbered sub-steps: "
            "(a) Formwork, (b) Reinforcement, (c) Mix Design / Batching, "
            "(d) Mixing, (e) Transportation & Placing, (f) Compaction, "
            "(g) Finishing, (h) Curing.  "
            "Include specific values (grades, w/c ratios, slump, curing period) "
            "from the specification where available."
        ),
    },
    {
        "key": "equipment",
        "heading": "6. Equipment Used",
        "queries": [
            "equipment machinery plant concrete mixer batching plant transit mixer pump",
            "vibrator needle poker vibration compaction equipment tools",
            "formwork shuttering props scaffolding equipment",
        ],
        "word_limit": 200,
        "instruction": (
            "List all plant, equipment, and tools mentioned in the specification "
            "for RCC work as a bullet list with a brief note on its use."
        ),
    },
    {
        "key": "personnel",
        "heading": "7. Key People Involved",
        "queries": [
            "personnel responsible engineer supervisor quality control inspection",
            "testing laboratory inspector site engineer foreman concrete",
        ],
        "word_limit": 150,
        "instruction": (
            "List roles/personnel mentioned or implied in the specification "
            "(e.g., site engineer, quality inspector, lab technician) and their "
            "responsibilities."
        ),
    },
    {
        "key": "quality",
        "heading": "8. Quality Control and Testing",
        "queries": [
            "quality control testing cube strength acceptance criteria concrete",
            "sampling frequency cube mould testing compressive strength",
            "inspection checks hold points concrete placement",
        ],
        "word_limit": 250,
        "instruction": (
            "Summarise quality-control requirements: sampling frequency, test "
            "types, acceptance criteria, and any hold/witness points from the "
            "specification."
        ),
    },
    {
        "key": "health_safety",
        "heading": "9. Health, Safety & Environment (HSE) Considerations",
        "queries": [
            "safety health environment precautions hazards concrete construction",
            "PPE protective equipment safety measures workers",
        ],
        "word_limit": 150,
        "instruction": (
            "Summarise HSE measures stated in the specification.  If none are "
            "explicitly stated, note that only and do NOT invent content."
        ),
    },
    {
        "key": "other",
        "heading": "10. Other Relevant Information",
        "queries": [
            "special requirements restrictions notes concrete specification",
            "rejection defective concrete remedial action non-conformance",
            "hot weather cold weather concreting special conditions",
        ],
        "word_limit": 150,
        "instruction": (
            "Include any other specification requirements not covered in the "
            "sections above (e.g. weather conditions, hot/cold weather "
            "concreting, repair of defective work)."
        ),
    },
]


# ══════════════════════════════════════════════════════════════════════════
# 2.  RETRIEVAL HELPERS
# ══════════════════════════════════════════════════════════════════════════

def load_vector_store(db_dir: str):
    """Open the persisted ChromaDB and return the collection."""
    client = chromadb.PersistentClient(path=db_dir)
    collection = client.get_or_create_collection("sections")
    return collection


def _query_collection(collection, embedding: list,
                      n_results: int,
                      chunk_type_filter: str | None = None) -> list[tuple]:
    """
    Single ChromaDB query.  Returns list of (id, doc, meta, distance).
    Optionally filters to a specific chunk_type ("section" or "table").
    Drops results whose cosine distance exceeds the relevance threshold.
    """
    DISTANCE_THRESHOLD = 1.2   # cosine distance: lower = more similar (0 = identical)

    kwargs = dict(
        query_embeddings=embedding,
        n_results=n_results,
        include=["documents", "metadatas", "distances"],
    )
    if chunk_type_filter:
        kwargs["where"] = {"chunk_type": chunk_type_filter}

    try:
        res = collection.query(**kwargs)
    except Exception:
        # ChromaDB raises if the filtered subset is smaller than n_results
        kwargs["n_results"] = 1
        try:
            res = collection.query(**kwargs)
        except Exception:
            return []

    rows = []
    for doc, meta, cid, dist in zip(
        res["documents"][0],
        res["metadatas"][0],
        res["ids"][0],
        res["distances"][0],
    ):
        if dist <= DISTANCE_THRESHOLD:
            rows.append((cid, doc, meta, dist))
    return rows



def retrieve_chunks(collection, model: SentenceTransformer,
                    queries: list[str], n_per_query: int = 5) -> list[dict]:

    seen_ids: set[str] = set()
    raw: list[tuple] = []   # (id, doc, meta, distance)

    for query in queries:
        embedding = model.encode([query]).tolist()

        # Pass 1 — prose + table (unfiltered)
        for row in _query_collection(collection, embedding, n_per_query):
            if row[0] not in seen_ids:
                seen_ids.add(row[0])
                raw.append(row)

        # Pass 2 — tables only
        for row in _query_collection(collection, embedding,
                                     max(2, n_per_query // 2),
                                     chunk_type_filter="table"):
            if row[0] not in seen_ids:
                seen_ids.add(row[0])
                raw.append(row)

    # ── Reranking ─────────────────────────────────────────────
    reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

    combined_query = " ".join(queries)
    pairs = [(combined_query, row[1]) for row in raw]

    scores = reranker.predict(pairs)

    # Attach scores
    raw = [(*row, scores[i]) for i, row in enumerate(raw)]

    # Sort by reranker score (descending)
    raw.sort(key=lambda r: r[4], reverse=True)

    # Optional: keep top-k after reranking
    raw = raw[:10]

    # ── Final output ───────────────────────────────────────────
    return [
        {
            "id":         cid,
            "section":    meta.get("section", ""),
            "title":      meta.get("title", ""),
            "chunk_type": meta.get("chunk_type", "section"),
            "table_id":   meta.get("table_id", ""),
            "text":       doc,
            "distance":   round(dist, 4),
        }
        for cid, doc, meta, dist, _ in raw
    ]


def pipe_to_markdown_table(text: str) -> str:
    """
    Convert the pipe-delimited table text produced by embed_and_store into a
    proper markdown table the LLM can reason over accurately.

    Input (first line is caption, rest are data rows):
        Table 4.1 Graded Stone Aggregate
        IS Sieve Designation | 40 mm | 20 mm | 16 mm
        80 mm | 100 | - | -

    Output:
        **Table 4.1 Graded Stone Aggregate**
        | IS Sieve Designation | 40 mm | 20 mm | 16 mm |
        |---|---|---|---|
        | 80 mm | 100 | - | - |
    """
    lines = [l for l in text.strip().splitlines() if l.strip()]
    if not lines:
        return text

    out = []
    data_lines = []
    caption    = ""

    # First line: check if it's a caption (no pipe) or a data row
    if "|" not in lines[0]:
        caption    = lines[0].strip()
        data_lines = lines[1:]
    else:
        data_lines = lines

    if caption:
        out.append(f"**{caption}**")

    if not data_lines:
        return "\n".join(out)

    # Build markdown table
    def _row(cells):
        return "| " + " | ".join(c.strip() for c in cells) + " |"

    header_cells = data_lines[0].split("|")
    out.append(_row(header_cells))
    out.append("| " + " | ".join("---" for _ in header_cells) + " |")  # separator

    for row in data_lines[1:]:
        out.append(_row(row.split("|")))

    return "\n".join(out)


def chunks_to_context(chunks: list[dict], max_chars: int = 10000) -> str:
    """
    Serialise retrieved chunks into a context string for the LLM.

    Improvements over v1:
    - Prose chunks and table chunks are formatted differently
    - Table chunks are rendered as markdown tables (not raw pipe text)
    - Prose and table sections are grouped with clear headers so the LLM
      knows what kind of content it is reading
    - Hard truncation replaced by per-chunk budget check so the most
      relevant chunks (sorted by distance) always appear first
    """
    prose_blocks = []
    table_blocks = []
    total_chars  = 0

    for i, c in enumerate(chunks, 1):
        chunk_type = c.get("chunk_type", "section")
        header     = f"[{i}] §{c['section']} — {c['title']}"

        if chunk_type == "table":
            table_label = c.get("table_id", "")
            body = pipe_to_markdown_table(c["text"])
            block = f"{header}\n{body}\n"
        else:
            block = f"{header}\n{c['text'].strip()}\n"

        if total_chars + len(block) > max_chars:
            break

        if chunk_type == "table":
            table_blocks.append(block)
        else:
            prose_blocks.append(block)

        total_chars += len(block)

    parts = []
    if prose_blocks:
        parts.append("=== SPECIFICATION TEXT ===\n" + "\n".join(prose_blocks))
    if table_blocks:
        parts.append("=== SPECIFICATION TABLES ===\n" + "\n".join(table_blocks))

    return "\n\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════
# 3.  LLM HELPERS  (Gemini 1.5 Flash)
# ══════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """\
You are a technical writer creating a Method Statement for Reinforced Cement \
Concrete (RCC) works on a construction project.

RULES:
1. Use ONLY information from the specification excerpts provided.
2. If the excerpts do not contain enough information for a section, write:
   "Not explicitly stated in the specification."
3. Do NOT hallucinate values, codes, or requirements.
4. Be concise and professional.
5. Cite the specification section number (e.g. "per §4.1.2") wherever possible.
6. The context contains two blocks:
   - SPECIFICATION TEXT  : prose paragraphs from the spec
   - SPECIFICATION TABLES: data tables in markdown format
   When referencing table values (sieve sizes, mix ratios, % passing etc.)
   read the column header and row label carefully before quoting a number.
"""


def call_gemini(api_key: str, context: str, section_instruction: str,
                word_limit: int, model_name: str = "gemini-2.5-flash") -> str:
    """Call Gemini and return the generated text for one MS section."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)

    prompt = (
        f"{SYSTEM_PROMPT}\n\n"
        f"--- SPECIFICATION EXCERPTS ---\n{context}\n\n"
        f"--- TASK ---\n{section_instruction}\n"
        f"Word limit: approximately {word_limit} words.\n"
        f"Write the section content now (no heading, just the body text):"
    )

    response = model.generate_content(prompt)
    return response.text.strip()


# ══════════════════════════════════════════════════════════════════════════
# 4.  WORD DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════

def build_word_document(sections_content: dict[str, str],
                        team_name: str,
                        out_path: str) -> None:
    """
    Build a properly formatted Method Statement Word document.
    sections_content: {section_key -> generated_text}
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Title page ────────────────────────────────────────────────────────
    title = doc.add_heading("METHOD STATEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading("Reinforced Cement Concrete (RCC) Works", level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: {team_name}\n").bold = True
    meta.add_run(f"Date: {datetime.today().strftime('%d %B %Y')}\n")
    meta.add_run(
        "Document Reference: Based on CPWD Prescriptive Specifications\n"
    )

    doc.add_page_break()

    # ── Sections ──────────────────────────────────────────────────────────
    for sec in MS_SECTIONS:
        key     = sec["key"]
        heading = sec["heading"]
        content = sections_content.get(key, "Content not generated.")

        doc.add_heading(heading, level=1)

        # Convert simple markdown-like lists the LLM sometimes emits
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # bullet items starting with -, *, •
            if re.match(r"^[-*•]\s+", stripped):
                p = doc.add_paragraph(
                    stripped.lstrip("-*• "), style="List Bullet"
                )
            # numbered items
            elif re.match(r"^\d+[\.\)]\s+", stripped):
                p = doc.add_paragraph(
                    re.sub(r"^\d+[\.\)]\s+", "", stripped),
                    style="List Number",
                )
            else:
                p = doc.add_paragraph(stripped)
            p.runs[0].font.size = Pt(11) if p.runs else None

        doc.add_paragraph()  # spacer between sections

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_heading("Note on Sources", level=2)
    doc.add_paragraph(
        "All information in this Method Statement has been extracted from the "
        "CPWD Prescriptive Specifications document using an NLP-based retrieval "
        "pipeline.  Where specification text was insufficient, this has been "
        "explicitly noted.  No external sources were used unless stated otherwise."
    )

    doc.save(out_path)
    print(f"✅  Word document saved → {out_path}")


# ══════════════════════════════════════════════════════════════════════════
# 5.  ACCURACY METRIC HELPERS
# ══════════════════════════════════════════════════════════════════════════

def compute_retrieval_coverage(sections_content: dict[str, str],
                                retrieved_chunks: dict[str, list[dict]]) -> dict:
    """
    Coverage metric v2 — token overlap (Jaccard similarity).

    The old metric checked whether a chunk's *title* appeared verbatim in the
    generated text, which almost never matched and produced near-zero scores
    for every section.

    This version tokenises both the generated text and each retrieved chunk's
    body, then computes:

        Jaccard = |generated_tokens ∩ chunk_tokens| / |generated_tokens ∪ chunk_tokens|

    The section score is the *mean* Jaccard across its retrieved chunks.
    Score range: 0 (no word overlap) → 1 (identical vocabulary).

    Additional per-section diagnostics:
        n_chunks        – how many chunks were retrieved
        n_table_chunks  – how many were table chunks
        best_chunk      – section ID of the most-overlapping chunk
    """
    def tokenise(text: str) -> set[str]:
        # lowercase alphanum tokens, ignore stopwords and single chars
        STOPWORDS = {"the", "a", "an", "of", "in", "to", "and", "or",
                     "is", "are", "be", "as", "for", "with", "by", "on",
                     "at", "from", "it", "its", "not", "per", "shall"}
        tokens = set(re.findall(r"[a-z0-9]+", text.lower()))
        return tokens - STOPWORDS - {t for t in tokens if len(t) <= 1}

    def jaccard(set_a: set, set_b: set) -> float:
        if not set_a or not set_b:
            return 0.0
        inter = len(set_a & set_b)
        union = len(set_a | set_b)
        return round(inter / union, 4) if union else 0.0

    scores = {}
    for key, chunks in retrieved_chunks.items():
        gen_tokens = tokenise(sections_content.get(key, ""))
        if not chunks:
            scores[key] = {"score": None, "n_chunks": 0,
                           "n_table_chunks": 0, "best_chunk": None}
            continue

        jaccard_scores = []
        best_score, best_chunk = -1, None

        for c in chunks:
            chunk_tokens = tokenise(c["text"])
            j = jaccard(gen_tokens, chunk_tokens)
            jaccard_scores.append(j)
            if j > best_score:
                best_score = j
                best_chunk = c["section"]

        mean_score = round(sum(jaccard_scores) / len(jaccard_scores), 4)
        n_table    = sum(1 for c in chunks if c.get("chunk_type") == "table")

        scores[key] = {
            "score"        : mean_score,
            "n_chunks"     : len(chunks),
            "n_table_chunks": n_table,
            "best_chunk"   : best_chunk,
        }

    return scores


# ══════════════════════════════════════════════════════════════════════════
# 6.  MAIN PIPELINE
# ══════════════════════════════════════════════════════════════════════════

def run_pipeline(db_dir: str, api_key: str, out_dir: str,
                 team_name: str = "Your Team Name",
                 n_per_query: int = 5) -> None:

    os.makedirs(out_dir, exist_ok=True)

    # ── Load resources ────────────────────────────────────────────────────
    print("⏳  Loading embedding model …")
    embed_model = SentenceTransformer("all-MiniLM-L6-v2")

    print("⏳  Connecting to ChromaDB …")
    collection  = load_vector_store(db_dir)
    total_docs  = collection.count()
    print(f"    → {total_docs} chunks in store")

    # ── Per-section retrieval + generation ────────────────────────────────
    sections_content: dict[str, str]        = {}
    all_retrieved:    dict[str, list[dict]] = {}

    for sec in MS_SECTIONS:
        key     = sec["key"]
        heading = sec["heading"]
        print(f"\n📄  Processing: {heading}")

        # Stage 4 — retrieve
        chunks = retrieve_chunks(
            collection, embed_model, sec["queries"], n_per_query
        )
        all_retrieved[key] = chunks
        print(f"    → {len(chunks)} unique chunks retrieved")

        context = chunks_to_context(chunks)

        # Stage 5 — generate
        print(f"    → Calling Gemini …")
        try:
            text = call_gemini(
                api_key       = api_key,
                context       = context,
                section_instruction = sec["instruction"],
                word_limit    = sec["word_limit"],
            )
        except Exception as e:
            print(f"    ⚠  Gemini error: {e}")
            text = f"Generation failed: {e}"

        sections_content[key] = text
        print(f"    → {len(text.split())} words generated")

    # ── Build Word document ───────────────────────────────────────────────
    word_path = os.path.join(out_dir, "Method_Statement_RCC.docx")
    build_word_document(sections_content, team_name, word_path)

    # ── Save raw JSON (for debugging / accuracy checks) ───────────────────
    debug_path = os.path.join(out_dir, "generated_sections_debug.json")
    with open(debug_path, "w", encoding="utf-8") as f:
        json.dump(sections_content, f, ensure_ascii=False, indent=2)
    print(f"✅  Debug JSON saved → {debug_path}")

    # ── Accuracy metric ───────────────────────────────────────────────────
    scores = compute_retrieval_coverage(sections_content, all_retrieved)
    print("\n📊  Retrieval-coverage scores (Jaccard token overlap, 0→1):")
    for k, v in scores.items():
        if v["score"] is None:
            print(f"    {k:<15}: no chunks")
        else:
            print(f"    {k:<15}: {v['score']:.4f}  "
                  f"(chunks={v['n_chunks']}, tables={v['n_table_chunks']}, "
                  f"best=§{v['best_chunk']})")
    metric_path = os.path.join(out_dir, "accuracy_metrics.json")
    with open(metric_path, "w") as f:
        json.dump(scores, f, indent=2)
    print(f"✅  Metrics saved → {metric_path}")


# ══════════════════════════════════════════════════════════════════════════
# 7.  CLI ENTRY-POINT
# ══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Stage 4+5: Retrieve from ChromaDB and generate Method Statement via Gemini."
    )
    parser.add_argument("--db_dir",    default="chroma_db",
                        help="Path to persisted ChromaDB directory")
    parser.add_argument("--out_dir",   default="output",
                        help="Directory for output files")
    parser.add_argument("--api_key",   default=os.getenv("GEMINI_API_KEY", ""),
                        help="Gemini API key (or set GEMINI_API_KEY env var)")
    parser.add_argument("--team_name", default="Team Name / Unstop ID",
                        help="Team name for the title page")
    parser.add_argument("--n_results", type=int, default=5,
                        help="Number of chunks retrieved per query (default 5)")
    args = parser.parse_args()

    if not args.api_key:
        raise ValueError(
            "Gemini API key required.  Pass --api_key or set GEMINI_API_KEY."
        )

    run_pipeline(
        db_dir    = args.db_dir,
        api_key   = args.api_key,
        out_dir   = args.out_dir,
        team_name = args.team_name,
        n_per_query = args.n_results,
    )